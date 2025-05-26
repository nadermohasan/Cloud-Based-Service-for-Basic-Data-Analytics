from flask import Flask, redirect, request, render_template, send_file
import os, time
from docx import Document
import fitz  # PyMuPDF
from io import BytesIO
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from supabase import create_client

# Supabase configuration
SUPABASE_URL = "https://xijwkncibxssvmxqcwzz.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6InhpandrbmNpYnhzc3ZteHFjd3p6Iiwicm9sZSI6ImFub24iLCJpYXQiOjE3NDc4Mjc5NzUsImV4cCI6MjA2MzQwMzk3NX0.BIp5iNYKmoPe70dD-XHZaVYxoZ0lD8l1pKeVpFi-DXM"
BUCKET_NAME = "documents"

app = Flask(__name__)
supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

# التصنيفات المحددة مسبقًا
CLASSIFICATION_TREE = {
    "Technical": [
        "report",
        "data",
        "analysis",
        "system",
        "computer",
        "framework",
        "code",
        "write",
    ],
    "Information Technology": [
        "Programming",
        "Data Science",
        "Operating Systems",
        "Database",
        "AI",
        "Distributed Systems",
        "Software Engineering",
        "Software Development",
        "Developer",
        "Software",
        "Security",
        "Cloud",
    ],
    "Other": [""],
}

def get_all_files():
    """Get list of all files from Supabase storage"""
    res = supabase.storage.from_(BUCKET_NAME).list()
    return [file['name'] for file in res]

def get_file_bytes(filename):
    """Get file bytes from Supabase storage"""
    res = supabase.storage.from_(BUCKET_NAME).download(filename)
    return BytesIO(res)
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# الصفحة الرئيسية
@app.route("/")
def home():
    return render_template("home.html")

# رفع الملفات
@app.route("/upload", methods=["GET", "POST"])
def upload_file():
    if request.method == "POST":
        if 'file' not in request.files:
            return redirect(request.url)
            
        file = request.files['file']
        if file.filename == '':
            return redirect(request.url)
            
        if file and allowed_file(file.filename):
            try:
                file.seek(0)
                file_data = file.read()
                
                supabase.storage.from_(BUCKET_NAME).upload(
                    file.filename,
                    file_data,
                    {"content-type": file.mimetype}
                )
                return render_template("upload.html", filename=file.filename)
                
            except Exception as e:
                return f"Upload failed! the file has already been uploaded or invalid filename {str(e)}", 500
                
    return render_template("upload.html")

# استخراج العنوان من metadata أو أول سطر من الملف
def extract_title(filename, file_bytes):
    try:
        if filename.lower().endswith(".pdf"):
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            meta = doc.metadata
            if meta and meta.get("title"):
                title = meta["title"].strip()
                if title:  # التأكد من أن العنوان ليس فارغاً
                    return title
            page = doc.load_page(0)
            text = page.get_text()
            return text.split("\n")[0].strip() if text else filename.rsplit(".", 1)[0]
        elif filename.lower().endswith(".docx"):
            doc = Document(file_bytes)
            title = doc.core_properties.title or doc.paragraphs[0].text.strip()
            return title if title else filename.rsplit(".", 1)[0]
    except:
        return filename.rsplit(".", 1)[0]  # استخدام اسم الملف كعنوان افتراضي

def extract_text(filename, file_bytes):
    try:
        if filename.lower().endswith(".pdf"):
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            return "\n".join(page.get_text() for page in doc) or ""  # إرجاع سلسلة فارغة بدلاً من None
        elif filename.lower().endswith(".docx"):
            doc = Document(file_bytes)
            return "\n".join(p.text for p in doc.paragraphs) or ""  # إرجاع سلسلة فارغة بدلاً من None
    except Exception as e:
        print(f"Error extracting text from {filename}: {str(e)}")
        return ""  # إرجاع سلسلة فارغة في حالة الخطأ


# فرز حسب العنوان
@app.route("/sort")
def sort_files():
    start = time.time()
    files = []
    for filename in get_all_files():
        file_bytes = get_file_bytes(filename)
        title = extract_title(filename, file_bytes)
        # التأكد من أن العنوان ليس None
        title = title or filename.rsplit(".", 1)[0]  # استخدام اسم الملف إذا كان العنوان None
        size = round(len(file_bytes.getbuffer()) / 1024, 2)
        files.append({"filename": filename, "title": title, "size": size})
    files.sort(key=lambda x: (x["title"] or "").lower())  # استخدام سلسلة فارغة إذا كان العنوان None
    elapsed = round(time.time() - start, 2)
    return render_template("sorted.html", documents=files, time=elapsed)

# البحث داخل الملفات
@app.route("/search", methods=["GET", "POST"])
def search():
    if request.method == "POST":
        keyword = request.form.get("keyword", "").strip().lower()
        if not keyword:
            return render_template("search.html", error="Please enter a search keyword")
            
        start = time.time()
        results = []
        
        for filename in get_all_files():
            try:
                file_bytes = get_file_bytes(filename)
                content = extract_text(filename, file_bytes).lower()  # لن تسبب خطأ لأن extract_text تعيد سلسلة فارغة الآن
                
                if keyword in content:
                    index = content.find(keyword)
                    snippet_start = max(0, index - 30)
                    snippet_end = index + 70
                    snippet = content[snippet_start: snippet_end].replace(
                        keyword, f"<mark>{keyword}</mark>"
                    )
                    results.append({
                        "filename": filename,
                        "snippet": snippet,
                        "keyword": keyword
                    })
                    
            except Exception as e:
                print(f"Error processing file {filename}: {str(e)}")
                continue  # تخطي الملف في حالة الخطأ
                
        elapsed = round(time.time() - start, 2)
        return render_template("search_results.html", 
                             results=results, 
                             keyword=keyword, 
                             time=elapsed,
                             count=len(results))
    
    return render_template("search.html")

# التصنيف
@app.route("/classify")
def classify_documents():
    start = time.time()
    classification = {"Technical": [], "Information Technology": [], "Other": []}

    for filename in get_all_files():
        try:
            file_bytes = get_file_bytes(filename)
            content = extract_text(filename, file_bytes)
            
            # التأكد من أن المحتوى ليس None واستخدام سلسلة فارغة بدلاً من ذلك
            content = content.lower() if content else ""
            
            scores = {"Technical": 0, "Information Technology": 0}

            for category in ["Technical", "Information Technology"]:
                for keyword in CLASSIFICATION_TREE[category]:
                    # تحويل الكلمة المفتاحية إلى حروف صغيرة مرة واحدة فقط
                    if keyword.lower() in content:
                        scores[category] += 1

            max_score = max(scores.values())
            if max_score == 0:
                best_category = "Other"
            else:
                # تحديد الفئة ذات الأعلى نقاط
                best_category = max(scores.items(), key=lambda x: x[1])[0]

            classification[best_category].append(filename)
            
        except Exception as e:
            print(f"Error classifying file {filename}: {str(e)}")
            classification["Other"].append(filename)
            continue

    elapsed = round(time.time() - start, 2)
    return render_template("classification.html", 
                         classifications=classification, 
                         time=elapsed)


# الإحصائيات
@app.route("/statistics")
def statistics():
    filenames = get_all_files()
    total_size = 0
    
    # حساب الحجم الكلي
    for name in filenames:
        try:
            file_bytes = get_file_bytes(name)
            total_size += len(file_bytes.getbuffer())
        except Exception as e:
            print(f"Error calculating size for {name}: {str(e)}")
            continue

    total_size = total_size / 1000  # KB

    # قياس وقت الفرز
    start = time.time()
    for name in filenames:
        try:
            title = extract_title(name, get_file_bytes(name))
            if title:  # مجرد تنفيذ العملية لقياس الوقت
                pass
        except:
            continue
    sort_time = round(time.time() - start, 2)

    # قياس وقت البحث
    start = time.time()
    sample_kw = "data"
    for name in filenames:
        try:
            content = extract_text(name, get_file_bytes(name))
            if content and sample_kw in content.lower():  # التحقق من وجود المحتوى أولاً
                pass
        except:
            continue
    search_time = round(time.time() - start, 2)

    # قياس وقت التصنيف
    start = time.time()
    for name in filenames:
        try:
            content = extract_text(name, get_file_bytes(name))
            if content:  # مجرد تنفيذ العملية لقياس الوقت
                pass
        except:
            continue
    classify_time = round(time.time() - start, 2)

    stats = {
        "num_files": len(filenames),
        "total_size": round(total_size, 2),
        "sort_time": sort_time,
        "search_time": search_time,
        "classify_time": classify_time,
        "avg_size": round(total_size / len(filenames), 2) if filenames else 0
    }
    
    return render_template("statistics.html", stats=stats)


@app.route("/download/<filename>")
def download_highlighted(filename):
    keyword = request.args.get("keyword", "").lower()
    file_bytes = get_file_bytes(filename)

    if filename.lower().endswith(".pdf"):
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text_instances = page.search_for(keyword, quads=False)
            for inst in text_instances:
                highlight = page.add_highlight_annot(inst)
                highlight.update()
        temp_file = BytesIO()
        doc.save(temp_file)
        doc.close()
        temp_file.seek(0)
        return send_file(temp_file, as_attachment=True, download_name=f"highlighted_{filename}", mimetype="application/pdf")

    elif filename.lower().endswith(".docx"):
        doc = Document(file_bytes)
        for para in doc.paragraphs:
            if keyword in para.text.lower():
                for run in para.runs:
                    if keyword in run.text.lower():
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        temp_docx = BytesIO()
        doc.save(temp_docx)
        temp_docx.seek(0)
        return send_file(temp_docx, as_attachment=True, download_name=f"highlighted_{filename}", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    return "Unsupported file format", 400

if __name__ == "__main__":
    app.run(debug=True)